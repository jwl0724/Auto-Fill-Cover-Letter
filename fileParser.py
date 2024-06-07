import docx
import re


REGEX_EXPRESSION = "\[%=.*?=%\]"


class NoMarkedFieldsError(Exception):
    # Raised when document does not have any marked fields to replace
    pass


class parser:
    def __init__(self, filename: str):
        try:
            self.document = docx.Document(f"testFiles/{filename}")
            self.marked_fields = {}
        except docx.opc.exceptions.PackageNotFoundError:
            raise FileNotFoundError(f"{filename} could not be found")
        except PermissionError:
            raise PermissionError(f"Could not access {filename} due to invalid permissions")

        # find marked fields on the doc    
        for paragraph in self.document.paragraphs:
            if not re.search(REGEX_EXPRESSION, paragraph.text):
                continue
            marked_fields_in_paragraph = re.findall(REGEX_EXPRESSION, paragraph.text)
            # add marked fields of paragraph into dictionary
            for marked_field in marked_fields_in_paragraph:
                self.marked_fields[marked_field] = "filler"
        if not self.marked_fields:
            raise NoMarkedFieldsError(f"{filename} has no marked fields to replace")


    def set_marked_field(self, marked_field: str, replacement_text: str) -> None:
        marked_field = marked_field.upper()
        if not re.match(marked_field, REGEX_EXPRESSION):
            marked_field = f"[%= {marked_field} =%]"
        if marked_field in self.marked_fields:
            self.marked_fields[marked_field] = replacement_text


    def replace_marked_fields(self) -> None:
        bold, font_size, italic, underline, font_family = None, None, None, None, None
        for marked_field in self.marked_fields:
            for paragraph in self.document.paragraphs:
                for run in paragraph.runs:
                    # get the most recent run formatting
                    if bold != run.bold: bold = run.bold
                    if font_family != run.font.name and run.font.name != None: font_family = run.font.name
                    if font_size != run.font.size: font_size = run.font.size
                    if underline != run.underline: underline = run.underline
                    if italic != run.italic: italic = run.italic

                    if (paragraph.text.find(marked_field) == -1):
                        continue

                    # apply most recent formatting
                    paragraph.text = paragraph.text.replace(marked_field, self.marked_fields[marked_field])
                    paragraph.style.font.size = font_size
                    paragraph.style.font.underline = underline
                    paragraph.style.font.bold = bold
                    paragraph.style.font.name = font_family
                    paragraph.style.font.italic = italic


    def save(self, filename: str) -> None:
        try:
            self.document.save(f"output/{filename}")
            print(f"{filename} successfully saved")
        except PermissionError:
            print("Permission error, perhaps the file is already opened?")